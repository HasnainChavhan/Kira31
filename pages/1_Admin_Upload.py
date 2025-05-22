import streamlit as st
import os
import docx
from PyPDF2 import PdfReader

st.set_page_config(page_title="Admin Document Upload", layout="centered")
st.title("🛠️ Admin Panel - Upload Documents")

UPLOAD_DIR = "uploaded_docs"
os.makedirs(UPLOAD_DIR, exist_ok=True)

def extract_text(file_path, file_type):
    if file_type == "text/plain":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    elif file_type == "application/pdf":
        with open(file_path, "rb") as f:
            reader = PdfReader(f)
            return "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    return ""

uploaded_files = st.file_uploader(
    "Upload one or more documents (TXT, PDF, DOCX)",
    type=["txt", "pdf", "docx"],
    accept_multiple_files=True
)

if uploaded_files:
    all_text = ""
    saved_files = []

    for file in uploaded_files:
        save_path = os.path.join(UPLOAD_DIR, file.name)
        with open(save_path, "wb") as f:
            f.write(file.getbuffer())
        saved_files.append(file.name)
        all_text += extract_text(save_path, file.type) + "\n\n"

    st.session_state.doc_content = all_text
    st.session_state.uploaded_files = saved_files
    st.success("✅ Document(s) uploaded and content saved.")

# Reload saved files on refresh
elif 'uploaded_files' in st.session_state:
    all_text = ""
    for file_name in st.session_state.uploaded_files:
        file_path = os.path.join(UPLOAD_DIR, file_name)
        if os.path.exists(file_path):
            file_type = "application/pdf" if file_name.endswith(".pdf") else (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if file_name.endswith(".docx") else "text/plain"
            )
            all_text += extract_text(file_path, file_type) + "\n\n"
    st.session_state.doc_content = all_text
    st.info("📄 Loaded previously uploaded documents.")
